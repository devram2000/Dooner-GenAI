from pdfminer.high_level import extract_text
import os
from tqdm import tqdm
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def clean_text(text):
    """
    Remove non-XML compatible characters from text.
    """
    return ''.join(char for char in text if char.isprintable())

# Initialize a Word Document
doc = Document()

# Folder containing your PDFs (in the current directory)
pdf_folder = './Sustainability Plans'

# List of PDF files
pdf_files = [f for f in os.listdir(pdf_folder) if f.endswith('.pdf')]

# Iterate over each PDF file in the folder with a progress bar
for filename in tqdm(pdf_files, desc='Processing PDFs', unit='pdf'):
    file_path = os.path.join(pdf_folder, filename)

    # Add "Document Title:" and the title (filename without extension) as a heading in the Word document
    title = os.path.splitext(filename)[0]
    doc.add_heading("Document Title: " + title, level=1)

    # Extract text from the PDF file using pdfminer.six
    text = extract_text(file_path)
    if text:  # Only add non-empty text
        cleaned_text = clean_text(text)
        paragraph = doc.add_paragraph(cleaned_text)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Add a page break after each PDF's text
    doc.add_page_break()

# Save the Word document
doc.save('corpus.docx')
